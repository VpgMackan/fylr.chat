"""Handler for docx files."""

from docx import Document
from io import BytesIO


def handle_docx(file_content: bytes, **kwargs) -> str:
    """
    Extract text from docx text files.

    Args:
        file_content: Raw file content as bytes
        **kwargs: Additional arguments (unused for text files)

    Returns:
        Extracted text as string
    """
    try:
        document = Document(BytesIO(file_content))
        full_text = []

        # Extract text from paragraphs
        for para in document.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)

        # Extract text from headers and footers
        for section in document.sections:
            for para in section.header.paragraphs:
                full_text.append(para.text)
            for para in section.footer.paragraphs:
                full_text.append(para.text)

        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from docx: {e}")
